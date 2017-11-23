#!/usr/bin/env python3
'''Say you have a text file of guest names. This guests.txt file has one
name per line, as follows:

Prof. Plum
Miss Scarlet
Col. Mustard
Al Sweigart
Robocop

Write a program that would generate a Word document with custom
invitations that look like Figure 13-11.

You can download
a sample guests.txt file from http://nostarch.com/automatestuff/.
'''

import docx

guestFile = open('guests.txt', 'r')
guestStr = guestFile.read()
guestLst = guestStr.split('\n')
guestFile.close()

doc = docx.Document('invitations_empty.docx')

for guest in guestLst:
    doc.add_paragraph('It would be a pleasure to have the company of',\
                      style='InvText')
    doc.add_paragraph(guest, style='InvName')
    doc.add_paragraph('At 111010 Memory Lane on the Evening of',\
                      style='InvText')
    doc.add_paragraph('April 1st', style='InvDate')
    doc.add_paragraph('at 7 o\'clock', style='InvText')
    doc.add_page_break()
    
doc.save('invitations.docx')
